"""Local document extraction and bounded text chunking."""

import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from pypdf import PdfReader

from .config import (
    DOCUMENT_CHUNK_CHARACTERS,
    DOCUMENT_CHUNK_OVERLAP,
    MAX_COMPRESSION_RATIO,
    MAX_DOCUMENT_CHUNKS,
    MAX_EXTRACTED_CHARACTERS,
    MAX_PDF_PAGES,
    MAX_UNCOMPRESSED_BYTES,
    UPLOAD_MAX_BYTES,
)

ALLOWED_EXTENSIONS = {
    ".txt",
    ".md",
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".csv",
    ".sql",
    ".xml",
    ".html",
    ".css",
    ".pdf",
    ".docx",
}


class DocumentExtractionError(ValueError):
    """Raised when an uploaded document cannot be safely extracted."""


def _check_archive_expansion(data: bytes) -> None:
    """Reject a ZIP-backed document before anything is decompressed.

    The upload limit bounds the compressed input only. A 199 KB DOCX can carry
    116 MB of XML, which costs seconds of CPU and hundreds of megabytes of
    memory to parse. The central directory declares the uncompressed size, so
    this can be refused without decompressing a single byte.
    """

    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            uncompressed = sum(info.file_size for info in archive.infolist())
    except zipfile.BadZipFile as exc:
        raise DocumentExtractionError(
            "The file is not a readable DOCX archive."
        ) from exc

    if uncompressed > MAX_UNCOMPRESSED_BYTES:
        raise DocumentExtractionError(
            f"The document expands to {uncompressed // (1024 * 1024)} MB, which "
            f"exceeds the {MAX_UNCOMPRESSED_BYTES // (1024 * 1024)} MB limit."
        )

    ratio = uncompressed / max(len(data), 1)
    if ratio > MAX_COMPRESSION_RATIO:
        raise DocumentExtractionError(
            f"The document has a compression ratio of {ratio:.0f}:1, above the "
            f"{MAX_COMPRESSION_RATIO:.0f}:1 limit. This is characteristic of a "
            "decompression bomb rather than a normal document."
        )


def _guard_length(parts: List[str], running: int, addition: str) -> int:
    """Accumulate extracted text, stopping at the character ceiling."""

    remaining = MAX_EXTRACTED_CHARACTERS - running
    if remaining <= 0:
        return running
    parts.append(addition[:remaining])
    return running + min(len(addition), remaining)


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))

    if reader.is_encrypted:
        raise DocumentExtractionError(
            "The PDF is encrypted. Remove the password and upload it again."
        )

    page_count = len(reader.pages)
    if page_count > MAX_PDF_PAGES:
        raise DocumentExtractionError(
            f"The PDF has {page_count} pages, above the {MAX_PDF_PAGES} page limit."
        )

    pages: List[str] = []
    total = 0
    for index, page in enumerate(reader.pages, start=1):
        if total >= MAX_EXTRACTED_CHARACTERS:
            break
        text = (page.extract_text() or "").strip()
        if text:
            total = _guard_length(pages, total, f"[Page {index}]\n{text}")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    _check_archive_expansion(data)

    document = Document(BytesIO(data))

    blocks: List[str] = []
    total = 0
    for paragraph in document.paragraphs:
        if total >= MAX_EXTRACTED_CHARACTERS:
            break
        text = paragraph.text.strip()
        if text:
            total = _guard_length(blocks, total, text)

    for table_index, table in enumerate(document.tables, start=1):
        if total >= MAX_EXTRACTED_CHARACTERS:
            break
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            total = _guard_length(
                blocks, total, f"[Table {table_index}]\n" + "\n".join(rows)
            )

    return "\n\n".join(block for block in blocks if block)


def _extract_text(data: bytes) -> str:
    if b"\x00" in data:
        raise DocumentExtractionError(
            "The file appears to be binary and is not a supported document."
        )
    return data.decode("utf-8-sig", errors="replace")[:MAX_EXTRACTED_CHARACTERS]


def chunk_text(
    text: str,
    chunk_characters: int = DOCUMENT_CHUNK_CHARACTERS,
    overlap: int = DOCUMENT_CHUNK_OVERLAP,
    max_chunks: int = MAX_DOCUMENT_CHUNKS,
) -> tuple[List[str], bool]:
    """Split text on nearby line boundaries with bounded overlap."""

    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        return [], False

    chunks: List[str] = []
    start = 0
    truncated = False

    while start < len(normalized):
        if len(chunks) >= max_chunks:
            truncated = True
            break

        hard_end = min(start + chunk_characters, len(normalized))
        end = hard_end
        if hard_end < len(normalized):
            newline = normalized.rfind("\n", start + chunk_characters // 2, hard_end)
            if newline > start:
                end = newline

        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= len(normalized):
            break
        start = max(end - overlap, start + 1)

    return chunks, truncated


def extract_document(filename: str, data: bytes) -> Dict[str, Any]:
    """Extract supported local text without retaining the original bytes."""

    safe_name = Path(filename or "document").name
    extension = Path(safe_name).suffix.lower()

    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentExtractionError(
            "Unsupported file type. Allowed extensions: "
            + ", ".join(sorted(ALLOWED_EXTENSIONS))
        )

    if not data:
        raise DocumentExtractionError("The uploaded file is empty.")

    if len(data) > UPLOAD_MAX_BYTES:
        raise DocumentExtractionError(
            f"The file exceeds the {UPLOAD_MAX_BYTES} byte upload limit."
        )

    try:
        if extension == ".pdf":
            text = _extract_pdf(data)
        elif extension == ".docx":
            text = _extract_docx(data)
        else:
            text = _extract_text(data)
    except DocumentExtractionError:
        raise
    except Exception as exc:
        raise DocumentExtractionError(
            f"Unable to extract text from {safe_name}: {type(exc).__name__}."
        ) from exc

    if not text.strip():
        raise DocumentExtractionError(
            "No extractable text was found in the uploaded document."
        )

    chunks, truncated = chunk_text(text)
    return {
        "filename": safe_name,
        "extension": extension,
        "size_bytes": len(data),
        "text": text.strip(),
        "chunks": chunks,
        "chunk_count": len(chunks),
        "truncated": truncated,
        "character_count": len(text.strip()),
        "estimated_tokens": max(1, (len(text.strip()) + 3) // 4),
    }
